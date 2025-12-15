"""DOCX Resume Parser"""
import io
from typing import Union
from docx import Document


class DOCXParser:
    """Parse DOCX files and extract text content"""
    
    @staticmethod
    def extract_text(file_content: Union[bytes, io.BytesIO]) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_content: DOCX file as bytes or BytesIO object
            
        Returns:
            Extracted text as string
        """
        if isinstance(file_content, bytes):
            file_content = io.BytesIO(file_content)
        
        try:
            doc = Document(file_content)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n".join(text)
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = " ".join(text.split())
        return text.strip()
